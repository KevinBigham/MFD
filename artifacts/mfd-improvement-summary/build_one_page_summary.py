from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "MFD_Improvement_Plan_One_Page.docx"

INK = RGBColor(24, 31, 42)
MUTED = RGBColor(85, 92, 105)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GREEN = RGBColor(36, 103, 71)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)
BORDER = "DADCE0"
FILL_BLUE = "E8EEF5"
FILL_GREEN = "EAF4EE"
FILL_GOLD = "FFF4D8"
FILL_RED = "FCE8E6"
FILL_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        child = mar.find(qn(f"w:{name}"))
        if child is None:
            child = OxmlElement(f"w:{name}")
            mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, size=9, bold=False, color=INK):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return run


def add_para(doc, text="", size=9, bold=False, color=INK, before=0, after=3, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_cell_title(cell, text, color=INK):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text.upper(), size=8.5, bold=True, color=color)


def add_cell_bullet(cell, label, detail, color=INK):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, f"{label}: ", size=7.7, bold=True, color=color)
    add_run(p, detail, size=7.7, color=INK)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9)
    normal.font.color.rgb = INK

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "Mr. Football Dynasty Improvement Plan", size=17, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    add_run(subtitle, "Plain-English one-page status brief for the next Codex run", size=9, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(callout, [7.35])
    c = callout.cell(0, 0)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(c, FILL_BLUE)
    set_cell_border(c, "B7C9DD")
    set_cell_margins(c, top=80, bottom=80, start=130, end=130)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Bottom line: ", size=9, bold=True, color=DARK_BLUE)
    add_run(
        p,
        "MFD already has a deep, playable dynasty foundation. The best next work is not a giant rewrite; it is shipping small, tested slices that make existing systems clearer, safer, and more exciting.",
        size=8.8,
        color=INK,
    )

    add_para(doc, "", size=1, after=2)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [2.43, 2.43, 2.43])
    fills = [FILL_GREEN, FILL_GOLD, FILL_RED]
    titles = [
        ("Ready to build on", GREEN),
        ("Needs work", GOLD),
        ("Watch the risks", RED),
    ]
    bullets = [
        [
            ("Core game", "deterministic sim, week advance, saves, setup, roster, draft, trades, cap, free agency."),
            ("World depth", "Chip, media, weather, awards, HOF, records, legacy, player profiles, rivalries."),
            ("Guardrails", "many tests already protect routes, save version, RNG, content, CI, and app shell behavior."),
            ("Best pattern", "existing forecast helpers show how to explain decisions before committing them."),
        ],
        [
            ("Clarity", "some screens know the data but do not explain what changed or what to do next."),
            ("Unwired helpers", "franchise tags, coach development, poaching, rivalry sidecar, and fast-lane setup need real call sites."),
            ("Mismatch bugs", "endorsements, scenario blockers, cap rule display, achievements, and waivers have source-truth gaps."),
            ("Living world", "player memory, offseason calendar, team flavor, and dynasty receipts can be stronger."),
        ],
        [
            ("Saves", "new persistent fields need schema, migration, defaults, fixtures, and old-save tests."),
            ("Sim math", "formula changes need samples, deterministic checks, and sanity ranges."),
            ("Sidecars", "browser-local archives are not automatically in portable dynasty cartridges."),
            ("Claims", "do not say a feature is live unless source shows a production store/lifecycle path."),
        ],
    ]

    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, fills[idx])
        set_cell_border(cell)
        set_cell_margins(cell, top=85, bottom=85, start=105, end=105)
        add_cell_title(cell, titles[idx][0], titles[idx][1])
        for label, detail in bullets[idx]:
            add_cell_bullet(cell, label, detail, titles[idx][1])

    add_para(doc, "", size=1, after=2)

    h = add_para(doc, "Future /goal plan", size=10.5, bold=True, color=BLUE, before=1, after=2)
    h.paragraph_format.keep_with_next = True
    plan = doc.add_table(rows=1, cols=4)
    plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    plan.autofit = False
    set_table_width(plan, [1.77, 1.88, 1.85, 1.85])
    plan_cells = plan.rows[0].cells
    plan_items = [
        ("1. Pick", "Score 3-5 slices by impact, risk, testability, and rollback."),
        ("2. Ship", "Finish one vertical slice with source, tests, docs, and verification."),
        ("3. Record", "Write a run ledger: files, tests, docs, risks, rollback, next slice."),
        ("4. Continue", "Only move on when the current slice is clean and independent."),
    ]
    for cell, (label, detail) in zip(plan_cells, plan_items):
        set_cell_shading(cell, FILL_GRAY)
        set_cell_border(cell)
        set_cell_margins(cell, top=80, bottom=80, start=100, end=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_run(p, label, size=8.4, bold=True, color=DARK_BLUE)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        add_run(p2, detail, size=7.5, color=INK)

    add_para(doc, "", size=1, after=2)

    h2 = add_para(doc, "Best next slices", size=10.5, bold=True, color=BLUE, before=1, after=1)
    h2.paragraph_format.keep_with_next = True
    priorities = [
        "Waiver/practice-squad clarity",
        "Achievement and player-memory clarity",
        "Endorsement source-of-truth fix",
        "Offseason calendar checklist",
        "Cap rule-awareness display",
        "Scenario constraint integrity",
        "Team-ops receipts",
        "Weather/game-day clarity",
    ]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.02
    add_run(p, "Start with read-only clarity when possible. Good candidates: ", size=8.2, color=INK)
    add_run(p, "; ".join(priorities) + ".", size=8.2, bold=True, color=INK)

    footer = doc.add_table(rows=1, cols=1)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.autofit = False
    set_table_width(footer, [7.35])
    f = footer.cell(0, 0)
    set_cell_shading(f, "FFFFFF")
    set_cell_border(f, "FFFFFF", size="0")
    set_cell_margins(f, top=30, bottom=20, start=0, end=0)
    fp = f.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    add_run(
        fp,
        "Simple rule for the marathon: improve what players feel, preserve saves and deterministic sim, and update the guide after every verified slice.",
        size=7.6,
        bold=True,
        color=MUTED,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
